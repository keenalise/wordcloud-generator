"""
Multi-format Word Cloud Generator
Supports: .odt, .docx, .pdf, .txt (and other plain-text formats)

Install dependencies:
    pip install wordcloud matplotlib nltk scikit-learn odfpy python-docx pypdf
"""

import os
import string
import matplotlib.pyplot as plt
from wordcloud import WordCloud
import nltk
from nltk.corpus import stopwords
from sklearn.feature_extraction.text import CountVectorizer

try:
    stopwords.words('english')
except LookupError:
    nltk.download('stopwords')

document_path = '/home/keen-alise/Documents/keywords/manoj sir.odt'
output_image_path = 'keyword_wordcloud.png'

color_map = {
    'ml_ai': '#1e3f71',
    'sports': '#2c7bb6',
    'methodology': '#0e5680',
    'technical_tools': '#5c9fc9'
}

color_groups = {
    'ml_ai': ['machine learning', 'prediction', 'betting model', 'data analysis',
              'logistic', 'accuracy', 'feature engineering', 'validation'],
    'sports': ['football', 'premier league', 'la liga', 'serie a', 'bundesliga',
               'ligue 1', 'value betting', 'odds'],
    'methodology': ['statistical inference', 'probability', 'regression',
                     'poisson distribution', 'classification', 'random forest'],
    'technical_tools': ['python']
}


def get_color_func(word, *args, **kwargs):
    word_lower = word.lower()
    for group_name, terms in color_groups.items():
        if word_lower in terms:
            return color_map[group_name]
    return '#36454f'


def extract_text_from_odt(file_path):
    from odf import text as odf_text, teletype
    from odf.opendocument import load

    doc = load(file_path)
    paragraphs = doc.getElementsByType(odf_text.P)
    return "\n".join(teletype.extractText(p) for p in paragraphs)


def extract_text_from_docx(file_path):
    import docx

    doc = docx.Document(file_path)
    parts = [p.text for p in doc.paragraphs]

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)

    return "\n".join(parts)


def extract_text_from_pdf(file_path):
    from pypdf import PdfReader

    reader = PdfReader(file_path)
    parts = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(parts)


def extract_text_from_plain(file_path):
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()


EXTRACTORS = {
    '.odt': extract_text_from_odt,
    '.docx': extract_text_from_docx,
    '.pdf': extract_text_from_pdf,
}


def extract_and_preprocess(file_path):
    if not os.path.isfile(file_path):
        print(f"Error: File not found at {file_path}. Please provide a valid path.")
        exit()

    ext = os.path.splitext(file_path)[1].lower()
    extractor = EXTRACTORS.get(ext, extract_text_from_plain)

    try:
        raw_text = extractor(file_path)
    except Exception as e:
        print(f"Error extracting text from '{file_path}' as '{ext}': {e}")
        exit()

    if not raw_text or not raw_text.strip():
        print("Warning: No text could be extracted from this file.")
        exit()

    text_lower = raw_text.lower()
    text_lower = text_lower.translate(str.maketrans('', '', string.punctuation))
    text_lower = text_lower.translate(str.maketrans('', '', string.digits))

    return text_lower


def extract_keywords(text, n_gram_range=(1, 2)):
    stop_words = set(stopwords.words('english'))
    stop_words.update(['figure', 'visual', 'representation', 'key', 'concepts',
                        'size', 'indicates', 'relative', 'importance', 'research'])

    vectorizer = CountVectorizer(
        ngram_range=n_gram_range,
        stop_words=list(stop_words),
        token_pattern=r'\b[^\d\W]+\b'
    )

    try:
        X = vectorizer.fit_transform([text])
    except ValueError:
        print("Error: Text is too short to process.")
        exit()

    all_ngrams = vectorizer.get_feature_names_out()
    counts = X.toarray().sum(axis=0)
    keyword_counts = dict(zip(all_ngrams, counts))

    all_allowed_terms = set()
    for group in color_groups.values():
        all_allowed_terms.update(group)

    filtered_keywords = {k: v for k, v in keyword_counts.items() if k in all_allowed_terms}

    if not filtered_keywords:
        print("No predefined keywords found in the document. Showing all extracted phrases instead.")
        filtered_keywords = keyword_counts

    return dict(sorted(filtered_keywords.items(), key=lambda item: item[1], reverse=True))


def generate_word_cloud(keyword_dict):
    wc = WordCloud(
        background_color="white",
        width=1200,
        height=800,
        max_words=200,
        prefer_horizontal=0.9,
        color_func=get_color_func,
        font_path=None
    )

    wc.generate_from_frequencies(keyword_dict)

    plt.figure(figsize=(15, 10))
    plt.imshow(wc, interpolation="bilinear")
    plt.axis("off")
    plt.title("Keywords", fontsize=24, pad=20, color=color_map['ml_ai'])

    plt.figtext(0.5, 0.02,
                "Figure 1: Visual representation of key concepts\n(Size indicates relative importance in the research)",
                ha="center", fontsize=12, style='italic', color='gray')

    print(f"Saving word cloud to {output_image_path}...")
    plt.savefig(output_image_path, bbox_inches='tight')
    plt.show()


if __name__ == '__main__':
    raw_text = extract_and_preprocess(document_path)
    keyword_frequencies = extract_keywords(raw_text)

    if keyword_frequencies:
        generate_word_cloud(keyword_frequencies)
        print("Success!")
    else:
        print("No keywords could be extracted. Check document.")
